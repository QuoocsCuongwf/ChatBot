"""
generate_report.py — Tạo báo cáo Word (.docx) chi tiết về dự án Legal RAG Chatbot
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime
import os

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professional styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    return table

def create_report():
    doc = Document()
    
    # ─── Page Setup ───
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    
    # ─── Styles ───
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    for i in range(1, 5):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # TRANG BÌA
    # ═══════════════════════════════════════════════════════════════════════════
    
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('BÁO CÁO DỰ ÁN')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('LEGAL RAG CHATBOT\nKiến Trúc 2 Tầng LLM')
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(51, 102, 153)
    
    doc.add_paragraph()
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run('Hệ thống hỏi đáp pháp luật Việt Nam\nsử dụng Retrieval-Augmented Generation (RAG)\nvới kiến trúc 2 tầng LLM để tối ưu chi phí và hiệu suất')
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(4):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y")}\nPhiên bản: 2.0 — 2-Tier Architecture')
    run.font.size = Pt(11)
    run.font.italic = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # MỤC LỤC
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        '1. Tổng quan dự án',
        '2. Kiến trúc hệ thống 2 Tầng',
        '   2.1. Sơ đồ pipeline tổng quát',
        '   2.2. Mô tả chi tiết từng module',
        '   2.3. Logic định tuyến 2 tầng (Tier Routing)',
        '3. Chi tiết các thành phần (Components)',
        '   3.1. Retriever (FAISS + Bi-encoder)',
        '   3.2. Cross-encoder Reranker',
        '   3.3. Context Builder',
        '   3.4. Gating Strategy',
        '   3.5. LLM Client (2 Tier)',
        '   3.6. Pipeline Logger',
        '   3.7. Evaluator',
        '4. Dữ liệu và Mô hình',
        '   4.1. FAISS Index',
        '   4.2. Cross-encoder (đã fine-tune)',
        '   4.3. Dữ liệu dev/train',
        '5. Kết quả đánh giá hiện tại',
        '   5.1. Kết quả test 10 mẫu',
        '   5.2. Phân bố Tier Routing',
        '   5.3. Nhật ký Excel (Pipeline Logger)',
        '6. Hướng triển khai mở rộng',
        '   6.1. Tầng 1 LOCAL — LLM local thực tế',
        '   6.2. Tầng 2 API — Gemini / GPT',
        '   6.3. Tích hợp Frontend (Web / App)',
        '   6.4. Scale dữ liệu pháp luật',
        '   6.5. Streaming & Real-time',
        '7. Tối ưu hóa (Optimization)',
        '   7.1. Tối ưu Retrieval',
        '   7.2. Tối ưu Gating Thresholds',
        '   7.3. Tối ưu LLM Generation',
        '   7.4. Tối ưu chi phí (Cost Optimization)',
        '   7.5. Tối ưu latency',
        '8. Các vấn đề đã giải quyết (Bug Fixes)',
        '9. Kế hoạch tiếp theo (Roadmap)',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        for run in p.runs:
            run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 1. TỔNG QUAN DỰ ÁN
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('1. Tổng quan dự án', level=1)
    
    doc.add_paragraph(
        'Dự án Legal RAG Chatbot là hệ thống hỏi đáp pháp luật Việt Nam sử dụng kỹ thuật '
        'Retrieval-Augmented Generation (RAG). Hệ thống được thiết kế theo kiến trúc 2 tầng LLM '
        '(2-Tier Architecture) nhằm tối ưu chi phí API, tốc độ phản hồi, và chất lượng câu trả lời.'
    )
    
    doc.add_heading('Mục tiêu hệ thống:', level=3)
    objectives = [
        'Trả lời câu hỏi pháp luật Việt Nam chính xác, kèm trích dẫn điều/khoản cụ thể',
        'Tự động phân loại mức độ phức tạp của câu hỏi để định tuyến đến LLM phù hợp',
        'Tiết kiệm chi phí API bằng cách dùng LLM local cho câu hỏi đơn giản',
        'Từ chối trả lời (abstain) khi không đủ căn cứ pháp lý — thay vì hallucinate',
        'Ghi log toàn bộ pipeline vào Excel để phân tích và cải thiện',
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('Công nghệ sử dụng:', level=3)
    add_styled_table(doc,
        ['Thành phần', 'Công nghệ', 'Mô tả'],
        [
            ['Bi-encoder', 'Quockhanh05/Vietnam_legal_embeddings', 'Mã hóa query + chunks thành vector (match FAISS index)'],
            ['Cross-encoder', 'Custom fine-tuned model', 'Rerank chunks, train trên dữ liệu FAISS-matched (loss 0.79→0.15)'],
            ['FAISS Index', 'faiss-cpu (cosine similarity)', '1861 vectors từ văn bản pháp luật VN (Nghị định, Thông tư...)'],
            ['LLM Tier 1', 'llama.cpp / Placeholder', 'LLM local, nhanh, miễn phí (score cao + margin lớn)'],
            ['LLM Tier 2', 'Gemini API / GPT', 'LLM API, chất lượng cao (score trung bình / cautious)'],
            ['Logging', 'openpyxl + CSV + JSONL', 'Nhật ký 3 sheet Excel: Chi tiết, Thống kê, Phân bố Tier'],
            ['Ngôn ngữ', 'Python 3.11', 'Conda environment: ocr311'],
            ['Framework', 'sentence-transformers, PyTorch', 'CUDA GPU cho encoding và inference'],
        ]
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 2. KIẾN TRÚC HỆ THỐNG 2 TẦNG
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('2. Kiến trúc hệ thống 2 Tầng', level=1)
    
    doc.add_heading('2.1. Sơ đồ pipeline tổng quát', level=2)
    
    doc.add_paragraph(
        'Pipeline xử lý một câu hỏi pháp luật đi qua 6 bước chính:'
    )
    
    pipeline_steps = [
        ['Bước', 'Module', 'Mô tả', 'Output'],
        ['1', 'Retrieve & Rerank', 'Bi-encoder tìm top-50 chunks từ FAISS → Cross-encoder rerank → top-10', 'List[ChunkInfo] sorted by score'],
        ['2', 'Context Building', 'Dedup (max 2 chunks/nguồn), Span Trim, Metadata Injection', 'context_string + processed_chunks'],
        ['3', 'Gating Decision', '5 rules: Score threshold, Margin, Keyword coverage, Legal keywords, Ambiguity', 'GatingDecision (ANSWER/CAUTIOUS/ABSTAIN/ASK_BACK)'],
        ['4', 'Tier Routing', 'Dựa trên score + margin → LOCAL / API / NONE', 'LLMTier (local/api/none)'],
        ['5', 'LLM Generation', 'Gọi LLM tầng tương ứng (hoặc skip nếu NONE)', 'RAGOutput với answer + citations'],
        ['6', 'Logging', 'Ghi 30 trường vào Excel/CSV/JSONL, phân tích tier', 'PipelineLogEntry → Excel 3 sheets'],
    ]
    
    table = doc.add_table(rows=len(pipeline_steps), cols=4)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(pipeline_steps):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('2.2. Mô tả chi tiết từng module', level=2)
    
    # Diagram as text
    p = doc.add_paragraph()
    run = p.add_run(
        '┌─────────────┐\n'
        '│   User      │\n'
        '│   Query     │\n'
        '└──────┬──────┘\n'
        '       ▼\n'
        '┌─────────────────┐     ┌──────────────┐\n'
        '│  Bi-encoder     │────►│  FAISS Index │\n'
        '│  (Legal HF)     │     │  (1861 vec)  │\n'
        '└──────┬──────────┘     └──────────────┘\n'
        '       ▼ top-50 candidates\n'
        '┌─────────────────┐\n'
        '│  Cross-encoder  │ Rerank → top-10\n'
        '│  (fine-tuned)   │\n'
        '└──────┬──────────┘\n'
        '       ▼\n'
        '┌─────────────────┐\n'
        '│ Context Builder │ Dedup + Trim + Metadata\n'
        '└──────┬──────────┘\n'
        '       ▼\n'
        '┌─────────────────┐\n'
        '│ Gating Strategy │ 5 Rules + Tier Routing\n'
        '└──────┬──────────┘\n'
        '       ▼\n'
        '  ┌────┴────┬──────────┐\n'
        '  ▼         ▼          ▼\n'
        '┌─────┐  ┌──────┐  ┌──────┐\n'
        '│LOCAL│  │ API  │  │ NONE │\n'
        '│Tier1│  │Tier2 │  │(skip)│\n'
        '└──┬──┘  └──┬───┘  └──┬───┘\n'
        '   ▼        ▼         ▼\n'
        '┌─────────────────────────┐\n'
        '│    Pipeline Logger      │\n'
        '│  (Excel + CSV + JSONL)  │\n'
        '└─────────────────────────┘\n'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('2.3. Logic định tuyến 2 tầng (Tier Routing)', level=2)
    
    doc.add_paragraph(
        'Sau khi Gating Strategy quyết định ANSWER hoặc CAUTIOUS, hệ thống xác định '
        'tầng LLM nào sẽ xử lý câu hỏi dựa trên 2 tiêu chí: score top-1 của cross-encoder '
        'và margin (chênh lệch giữa score top-1 và top-2).'
    )
    
    add_styled_table(doc,
        ['Tầng', 'Điều kiện', 'Khi nào sử dụng', 'Chi phí'],
        [
            ['Tier 1 — LOCAL', 'score ≥ 8.0 VÀ margin ≥ 0.3', 'Câu hỏi đơn giản, confidence cao, chunk top-1 vượt trội', 'Miễn phí (chạy local)'],
            ['Tier 2 — API', 'score ≥ -5.0 (nhưng không đủ LOCAL)', 'Câu hỏi phức tạp, cần diễn đạt chuẩn, CAUTIOUS decision', 'Tốn phí API (Gemini/GPT)'],
            ['NONE', 'score < -5.0 hoặc ABSTAIN/ASK_BACK', 'Không đủ căn cứ pháp lý → từ chối trả lời', 'Không tốn (skip LLM)'],
        ]
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Ngưỡng hiện tại (GatingConfig cho cross-encoder mode):'
    )
    
    add_styled_table(doc,
        ['Tham số', 'Giá trị', 'Ý nghĩa'],
        [
            ['threshold_pass', '3.0', 'Score ≥ 3.0 → ANSWER (cho phép generate)'],
            ['threshold_abstain', '-5.0', 'Score < -5.0 → ABSTAIN (từ chối)'],
            ['threshold_cautious', '0.0', 'Score 0.0 ~ 3.0 → CAUTIOUS (cảnh báo + route API)'],
            ['margin_min', '0.01', 'Margin tối thiểu (phân biệt top-1 và top-2)'],
            ['tier_local_min_score', '8.0', 'Score tối thiểu để dùng LOCAL LLM'],
            ['tier_local_min_margin', '0.3', 'Margin tối thiểu để dùng LOCAL LLM'],
            ['tier_api_min_score', '-5.0', 'Score tối thiểu để gọi API LLM'],
            ['min_keyword_coverage', '0.3', 'Overlap từ khóa tối thiểu 30%'],
            ['enable_ask_back', 'False', 'Tắt ask_back trong pipeline evaluation'],
        ]
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 3. CHI TIẾT CÁC THÀNH PHẦN
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('3. Chi tiết các thành phần (Components)', level=1)
    
    # --- 3.1 Retriever ---
    doc.add_heading('3.1. Retriever (FAISS + Bi-encoder)', level=2)
    doc.add_paragraph('File: cross-encoder/generation/run_generation.py — class LegalRetriever')
    doc.add_paragraph(
        'Retriever là bước đầu tiên trong pipeline, chịu trách nhiệm tìm kiếm '
        'các chunk văn bản pháp luật liên quan đến câu hỏi của người dùng.'
    )
    
    doc.add_heading('Quy trình:', level=3)
    steps_retriever = [
        'Encode query bằng bi-encoder (Quockhanh05/Vietnam_legal_embeddings) → vector 768 chiều',
        'Search FAISS index (cosine similarity) → top-50 candidates',
        'Dedup theo hash text[:200] để loại bỏ chunks trùng',
        'Extract metadata từ JSON: van_ban, dieu, khoan, diem, source_file',
        'Cross-encoder rerank top-50 → sắp xếp lại theo score → trả về top-10',
    ]
    for i, step in enumerate(steps_retriever, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')
    
    doc.add_heading('Cấu hình:', level=3)
    add_styled_table(doc,
        ['Tham số', 'Giá trị', 'Mô tả'],
        [
            ['top_k_retrieve', '50', 'Số candidates từ FAISS'],
            ['top_k_rerank', '10', 'Số chunks sau cross-encoder rerank'],
            ['normalize_embeddings', 'True', 'Chuẩn hóa vector cho cosine similarity'],
            ['batch_size (rerank)', '32', 'Batch size cho cross-encoder predict'],
            ['max_length (CE)', '256', 'Max token length cho cross-encoder input'],
        ]
    )
    
    # --- 3.2 Cross-encoder ---
    doc.add_heading('3.2. Cross-encoder Reranker', level=2)
    doc.add_paragraph('File: cross-encoder/outputs/models/cross_encoder_faiss_matched/saved_model/')
    doc.add_paragraph(
        'Cross-encoder đã được fine-tune trên dữ liệu FAISS-matched, đảm bảo model rerank '
        'chính xác trên chính dữ liệu trong FAISS index. '
        'Training loss giảm từ 0.79 xuống 0.15 sau 3 epochs.'
    )
    
    add_styled_table(doc,
        ['Thuộc tính', 'Giá trị'],
        [
            ['Base model', 'Mô hình phù hợp Legal Vietnamese'],
            ['Training data', 'FAISS metadata-matched pairs'],
            ['Loss function', 'CrossEntropyLoss'],
            ['Final loss', '0.15 (từ 0.79 ban đầu)'],
            ['Epochs', '3'],
            ['Score range', '~(-10 đến +12) — logit scores'],
            ['ROC-AUC', '~0.96 (trên tập dev)'],
        ]
    )
    
    # --- 3.3 Context Builder ---
    doc.add_heading('3.3. Context Builder', level=2)
    doc.add_paragraph('File: cross-encoder/generation/context_builder.py — class ContextBuilder (~458 dòng)')
    doc.add_paragraph(
        'Context Builder xử lý chunks đã rerank trước khi đưa vào prompt cho LLM. '
        '3 kỹ thuật chính:'
    )
    
    doc.add_heading('(a) Dedup & Diversify:', level=3)
    doc.add_paragraph(
        'Loại bỏ chunks trùng lặp và đa dạng hóa nguồn. '
        'Max 2 chunks từ cùng (văn_bản, điều). '
        'Sử dụng text hash[:200] để phát hiện duplicate. '
        'Kết quả: giảm redundancy, tăng coverage cho câu trả lời.'
    )
    
    doc.add_heading('(b) Span Trim:', level=3)
    doc.add_paragraph(
        'Cắt chunk chỉ giữ phần liên quan đến query. '
        'Tìm keyword matches → expand context_window=150 chars → merge overlapping spans. '
        'Giảm noise, giúp LLM focus vào thông tin quan trọng.'
    )
    
    doc.add_heading('(c) Metadata Injection:', level=3)
    doc.add_paragraph(
        'Inject header pháp lý vào mỗi chunk: [VB: Nghị định X, Điều Y, Khoản Z]. '
        'Giúp LLM biết nguồn trích dẫn chính xác để tạo citation.'
    )
    
    # --- 3.4 Gating Strategy ---
    doc.add_heading('3.4. Gating Strategy', level=2)
    doc.add_paragraph('File: cross-encoder/generation/gating.py — class GatingStrategy (~492 dòng)')
    doc.add_paragraph(
        'Gating Strategy là "bộ não" quyết định hệ thống có nên trả lời câu hỏi hay không, '
        'và nếu trả lời thì dùng tầng LLM nào. Gồm 5 rules tuần tự:'
    )
    
    rules = [
        ['Rule', 'Kiểm tra', 'Kết quả nếu fail', 'Tier'],
        ['Rule 1', 'Score threshold: top1_score < threshold_abstain (-5.0)?', 'ABSTAIN — confidence quá thấp', 'NONE'],
        ['Rule 2', 'Margin: (top1 - top2) < margin_min (0.01)?', 'ASK_BACK — kết quả không rõ ràng', 'NONE'],
        ['Rule 3', 'Keyword coverage < 30%?', 'ASK_BACK hoặc CAUTIOUS (route API)', 'NONE / API'],
        ['Rule 4', 'Legal keywords requirement check', 'ABSTAIN — thiếu từ khóa pháp lý cần thiết', 'NONE'],
        ['Rule 5', 'Query ambiguity check', 'ASK_BACK — câu hỏi cần làm rõ', 'NONE'],
        ['Final', 'Tất cả rules pass → Tier routing', 'ANSWER hoặc CAUTIOUS', 'LOCAL / API'],
    ]
    
    table = doc.add_table(rows=len(rules), cols=4)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(rules):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Phương thức _determine_tier(top1_score, margin) xác định tầng LLM:\n'
        '• score ≥ 8.0 VÀ margin ≥ 0.3 → LOCAL\n'
        '• score ≥ -5.0 (nhưng không đạt LOCAL) → API\n'
        '• Còn lại → NONE'
    )
    
    # --- 3.5 LLM Client ---
    doc.add_heading('3.5. LLM Client (2 Tier)', level=2)
    doc.add_paragraph('File: cross-encoder/generation/llm_client.py — class LLMClient (~621 dòng)')
    doc.add_paragraph(
        'LLM Client hỗ trợ 6 backends: llama_cpp, openrouter, openai, gemini, huggingface, '
        'placeholder. Hệ thống tạo 2 instance riêng biệt:'
    )
    
    add_styled_table(doc,
        ['Tier', 'Client', 'Config', 'max_tokens', 'context_length'],
        [
            ['Tier 1 (LOCAL)', 'local_client', 'llama_cpp / placeholder', '256', '2048'],
            ['Tier 2 (API)', 'api_client', 'gemini / openai', '512 (demo: 256)', '4096 (demo: 2048)'],
        ]
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Tính năng auto-detect: tự động phát hiện backend từ environment variables '
        '(GEMINI_API_KEY, OPENAI_API_KEY, OPENROUTER_API_KEY) hoặc local model (.gguf files).'
    )
    
    # --- 3.6 Pipeline Logger ---
    doc.add_heading('3.6. Pipeline Logger', level=2)
    doc.add_paragraph('File: cross-encoder/generation/pipeline_logger.py — class PipelineLogger (~580 dòng)')
    doc.add_paragraph(
        'Logger chuyên nghiệp ghi toàn bộ pipeline vào 3 format: Excel, CSV, JSONL. '
        'Mỗi query = 1 dòng log với 30 trường thông tin.'
    )
    
    doc.add_heading('30 trường log (PipelineLogEntry):', level=3)
    
    log_fields = [
        ['Nhóm', 'Trường', 'Ví dụ'],
        ['ID', 'timestamp, query_id, query', '2026-02-27 21:27:27, dev_0, "Điều 5 quy định gì?"'],
        ['Retrieval', 'num_chunks_retrieved, num_chunks_after_dedup, top1_chunk_id, top1_text_preview', '10, 5, 342, "Điều 5. Thẩm quyền..."'],
        ['Scores', 'score_retrieval_top1, score_rerank_top1, score_rerank_top2, margin', '0.85, 10.72, 10.16, 0.56'],
        ['Gating', 'gating_decision, gating_reason, gating_confidence', 'answer, "Passed all checks → Tier local", 10.72'],
        ['Tier', 'tier, tier_reason', 'local, "Score 10.7 >= 8.0, Margin 0.6 >= 0.3"'],
        ['LLM', 'llm_backend, llm_model, generated_answer, answer_length', 'placeholder, default, "[Placeholder]...", 156'],
        ['Citations', 'num_citations, citations_str, citation_hit', '1, "Điều 5, Khoản 2", True'],
        ['Latency', 'latency_retrieve_ms, latency_rerank_ms, latency_context_ms, latency_llm_ms, latency_total_ms', '1200, 800, 5, 150, 2200'],
        ['Status', 'error, status', '(trống), ok'],
    ]
    
    table = doc.add_table(rows=len(log_fields), cols=3)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(log_fields):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_heading('Excel 3 Sheets:', level=3)
    
    excel_sheets = [
        'Sheet 1 — "Chi tiết Pipeline": Toàn bộ 30 cột × N dòng, freeze panes ở D2, '
        'màu tier (xanh lá=LOCAL, xanh dương=API, đỏ nhạt=NONE), màu decision, auto-width cột.',
        'Sheet 2 — "Thống kê tổng hợp": Phân bố Tier, Phân bố Gating Decision, '
        'Thống kê Latency (avg/min/max), Tỷ lệ lỗi.',
        'Sheet 3 — "Phân bố Tier": Phân tích chi tiết mỗi tier (count, %, avg score, avg margin, '
        'avg latency, errors) + Phân tích chi phí (bao nhiêu % miễn phí, bao nhiêu % tốn API).',
    ]
    for sheet in excel_sheets:
        doc.add_paragraph(sheet, style='List Bullet')
    
    # --- 3.7 Evaluator ---
    doc.add_heading('3.7. Evaluator', level=2)
    doc.add_paragraph('File: cross-encoder/generation/evaluator.py — class GenerationEvaluator (~595 dòng)')
    doc.add_paragraph(
        'Evaluator đo lường chất lượng end-to-end của pipeline trên goldset (dev.jsonl). '
        'Metrics chính:'
    )
    
    add_styled_table(doc,
        ['Metric', 'Mô tả', 'Cách tính'],
        [
            ['Citation Hit Rate', 'Tỷ lệ câu trả lời trích dẫn đúng điều/khoản', 'Matched citations / Total'],
            ['Citation Precision', 'Độ chính xác trích dẫn', 'Correct cited / All cited'],
            ['Citation Recall', 'Độ phủ trích dẫn', 'Correct cited / Expected'],
            ['Citation F1', 'Harmonic mean P+R', '2*P*R / (P+R)'],
            ['Pass Rate', 'Tỷ lệ ANSWER decision', 'ANSWER count / Total'],
            ['Abstain Rate', 'Tỷ lệ từ chối trả lời', 'ABSTAIN count / Total'],
            ['Tier distribution', 'Phân bố LOCAL / API / NONE', 'Count per tier / Total'],
            ['API Savings', '% queries không gọi API', '(LOCAL + NONE) / Total'],
        ]
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 4. DỮ LIỆU VÀ MÔ HÌNH
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('4. Dữ liệu và Mô hình', level=1)
    
    doc.add_heading('4.1. FAISS Index', level=2)
    add_styled_table(doc,
        ['Thuộc tính', 'Giá trị'],
        [
            ['Đường dẫn', 'vector_data/legal_hf_cosine/'],
            ['Số vectors', '1861'],
            ['Loại index', 'IndexFlatIP (Inner Product = cosine similarity với normalized vectors)'],
            ['Metadata format', 'JSON array: [{text: "...", metadata: {van_ban, dieu, khoan, ...}}, ...]'],
            ['Nguồn dữ liệu', 'Nghị định, Thông tư, Nghị quyết pháp luật Việt Nam'],
            ['Embedding model', 'Quockhanh05/Vietnam_legal_embeddings (768 dim)'],
        ]
    )
    
    doc.add_heading('4.2. Cross-encoder (đã fine-tune)', level=2)
    add_styled_table(doc,
        ['Thuộc tính', 'Giá trị'],
        [
            ['Đường dẫn', 'cross-encoder/outputs/models/cross_encoder_faiss_matched/saved_model/'],
            ['Base model', 'Sentence-transformers CrossEncoder'],
            ['Training data', 'cross-encoder/data/train_faiss.jsonl (FAISS metadata-matched)'],
            ['Dev data', 'cross-encoder/data/dev_faiss.jsonl'],
            ['Positive samples', 'Query + chunk text từ FAISS metadata (label=1)'],
            ['Negative samples', 'Random negatives (label=0)'],
            ['Final training loss', '0.15'],
            ['Max input length', '256 tokens'],
        ]
    )
    
    doc.add_heading('4.3. Dữ liệu dev/train', level=2)
    add_styled_table(doc,
        ['File', 'Mô tả', 'Số lượng'],
        [
            ['cross-encoder/data/train_faiss.jsonl', 'Training data FAISS-matched (synthetic QA)', '~5000+ pairs'],
            ['cross-encoder/data/dev_faiss.jsonl', 'Dev/eval data', '~500+ pairs'],
            ['cross-encoder/data/train.jsonl', 'Original training data', 'Các cặp Q-A gốc'],
            ['cross-encoder/data/dev.jsonl', 'Dev data (dùng cho eval pipeline)', 'Các cặp Q-A dev'],
            ['retrieval/goldset.jsonl', 'Goldset retrieval đánh giá', 'Goldset chính'],
        ]
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 5. KẾT QUẢ ĐÁNH GIÁ
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('5. Kết quả đánh giá hiện tại', level=1)
    
    doc.add_heading('5.1. Kết quả test 10 mẫu', level=2)
    doc.add_paragraph(
        'Kết quả chạy: python cross-encoder/generation/run_generation.py --eval --max_samples 10 -v '
        '--local-backend placeholder --backend placeholder'
    )
    
    add_styled_table(doc,
        ['Metric', 'Giá trị', 'Ghi chú'],
        [
            ['Tổng samples', '10', 'Từ dev.jsonl (unique queries)'],
            ['Pass Rate', '90%', '9/10 queries được trả lời'],
            ['Abstain Rate', '10%', '1/10 bị abstain (NONE tier)'],
            ['Tier LOCAL', '1 (10%)', 'Score 10.72, Margin 0.56 → đủ điều kiện LOCAL'],
            ['Tier API', '8 (80%)', 'Score trung bình hoặc margin thấp'],
            ['Tier NONE', '1 (10%)', 'ABSTAIN (Rule 4: thiếu legal keywords trong context)'],
            ['API Savings', '20%', 'LOCAL(10%) + NONE(10%) = 20% queries không cần API'],
        ]
    )
    
    doc.add_heading('5.2. Phân bố Tier Routing', level=2)
    doc.add_paragraph(
        'Biểu đồ phân bố tier (từ logger output):'
    )
    
    p = doc.add_paragraph()
    run = p.add_run(
        '  ┌─ TIER ROUTING ─────────────────────────┐\n'
        '  │ T1 LOCAL (free)     1  10.0%  ███       │\n'
        '  │ T2 API (paid)       8  80.0%  ████████████████████████  │\n'
        '  │ NONE (skip)         1  10.0%  ███       │\n'
        '  │ API savings: 20.0%                      │\n'
        '  └─────────────────────────────────────────┘\n'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading('5.3. Nhật ký Excel (Pipeline Logger)', level=2)
    doc.add_paragraph(
        'Mỗi lần chạy tạo 3 file trong cross-encoder/generation/outputs/generation/logs/:'
    )
    
    add_styled_table(doc,
        ['File', 'Nội dung', 'Chi tiết'],
        [
            ['run_YYYYMMDD_HHMMSS.xlsx', 'Excel 3 sheets', '11 rows × 30 cols (Chi tiết), Stats, Tier analysis'],
            ['run_YYYYMMDD_HHMMSS.csv', 'CSV backup', 'Headers tiếng Việt (UTF-8-sig), dễ mở Excel'],
            ['run_YYYYMMDD_HHMMSS.jsonl', 'JSONL raw data', 'Structured data cho programmatic analysis'],
        ]
    )
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 6. HƯỚNG TRIỂN KHAI MỞ RỘNG
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('6. Hướng triển khai mở rộng', level=1)
    
    doc.add_heading('6.1. Tầng 1 LOCAL — LLM local thực tế', level=2)
    doc.add_paragraph(
        'Hiện tại Tier 1 đang dùng placeholder. Để triển khai thực tế:'
    )
    
    local_llm_options = [
        ['Phương án', 'Model', 'RAM cần', 'Tốc độ', 'Chất lượng'],
        ['llama.cpp (GGUF)', 'Vistral-7B-Q4 / Vinallama-7B-Q4', '~4-5 GB', 'Rất nhanh (~20 tok/s GPU)', 'Tốt cho câu hỏi đơn giản'],
        ['llama.cpp (GGUF)', 'Vistral-7B-Q8', '~8 GB', 'Nhanh (~15 tok/s GPU)', 'Tốt hơn Q4'],
        ['Ollama', 'llama3.1:8b / gemma2:9b', '~5-10 GB', 'Nhanh', 'Tốt, dễ cài đặt'],
        ['vLLM', 'Bất kỳ model HuggingFace', '~16+ GB GPU', 'Rất nhanh (batched)', 'Production-grade'],
        ['HuggingFace', 'ProtonX Legal Model (local)', '~14 GB', 'Trung bình', 'Đặc thù pháp luật VN'],
    ]
    
    table = doc.add_table(rows=len(local_llm_options), cols=5)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(local_llm_options):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Cách cài đặt nhanh Tier 1 LOCAL với llama.cpp:\n'
        '  1. Download model GGUF: https://huggingface.co/vilm/vistral-7b-chat-gguf\n'
        '  2. Build llama.cpp: cmake -B build && cmake --build build\n'
        '  3. Chạy: python run_generation.py --local-backend llama_cpp '
        '--local-model models/vistral-7b-q4.gguf --backend gemini --eval -v'
    )
    
    doc.add_heading('6.2. Tầng 2 API — Gemini / GPT', level=2)
    doc.add_paragraph(
        'Để kích hoạt Tier 2 API thực tế:'
    )
    
    api_steps = [
        'Set API key: $env:GEMINI_API_KEY = "your-key-here" (PowerShell)',
        'Cập nhật model name trong llm_client.py: gemini-pro → gemini-2.0-flash (hoặc gemini-1.5-flash)',
        'Chạy: python run_generation.py --backend gemini --eval --max_samples 50 -v',
        'Chi phí ước tính: Gemini Flash ~$0.075/1M input tokens → ~500 queries ≈ $0.01-0.05',
    ]
    for step in api_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('6.3. Tích hợp Frontend (Web / App)', level=2)
    doc.add_paragraph(
        'Hệ thống có thể tích hợp vào frontend thông qua REST API hoặc WebSocket:'
    )
    
    frontend_options = [
        ['Phương án', 'Công nghệ', 'Ưu điểm', 'Nhược điểm'],
        ['FastAPI + React', 'Python backend + React/Next.js', 'Dễ triển khai, streaming support', 'Cần deploy riêng'],
        ['Gradio', 'Gradio ChatInterface', 'Rất nhanh MVP, share link', 'UI hạn chế'],
        ['Streamlit', 'Streamlit Chat', 'Dễ code, UI đẹp', 'Không streaming tốt'],
        ['Flask + WebSocket', 'Flask-SocketIO', 'Real-time streaming', 'Phức tạp hơn'],
        ['Telegram Bot', 'python-telegram-bot', 'Tiếp cận user dễ', 'Giới hạn UI'],
    ]
    
    table = doc.add_table(rows=len(frontend_options), cols=4)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(frontend_options):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Khuyến nghị: FastAPI backend + React frontend cho production. '
        'Gradio cho demo/prototype nhanh.'
    )
    
    doc.add_heading('6.4. Scale dữ liệu pháp luật', level=2)
    doc.add_paragraph(
        'Hiện tại FAISS index chỉ có 1861 vectors. Để mở rộng:'
    )
    
    scale_steps = [
        'Thu thập thêm văn bản pháp luật: Nghị định mới, Luật, Pháp lệnh, Công văn hướng dẫn',
        'Chạy pipeline OCR (OCR_paddle_protonX.py) cho các file PDF/scan',
        'Parse cấu trúc: legal_parser.py → chunks.py → update_khoan.py',
        'Encode chunks mới: encode_legal_hf.py → thêm vào FAISS index',
        'Retrain cross-encoder trên dữ liệu mở rộng',
        'Cân nhắc chuyển từ IndexFlatIP sang IndexIVFFlat (>10K vectors) để tăng tốc search',
    ]
    for step in scale_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('6.5. Streaming & Real-time', level=2)
    doc.add_paragraph(
        'Hệ thống hiện chưa hỗ trợ streaming. Để thêm:'
    )
    streaming_items = [
        'LLM Client đã có enable_streaming flag trong LLMConfig — cần implement generator methods',
        'Dùng Server-Sent Events (SSE) hoặc WebSocket để stream tokens đến frontend',
        'Hiển thị tier và gating decision ngay trước khi LLM bắt đầu generate',
        'Streaming chỉ áp dụng cho Tier 2 API (Tier 1 local thường đủ nhanh)',
    ]
    for item in streaming_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 7. TỐI ƯU HÓA
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('7. Tối ưu hóa (Optimization)', level=1)
    
    doc.add_heading('7.1. Tối ưu Retrieval', level=2)
    
    retrieval_opts = [
        ['Vấn đề', 'Giải pháp', 'Ưu tiên'],
        ['FAISS index có ~1861 vec, một số trùng', 'Rebuild index sau khi dedup metadata', 'Cao'],
        ['top_k=50 có thể không đủ diverse', 'Tăng top_k_retrieve lên 100, giữ top_k_rerank=10', 'Trung bình'],
        ['Cosine similarity chỉ sử dụng dense', 'BM25 hybrid: kết hợp sparse (TF-IDF) + dense (FAISS)', 'Cao'],
        ['Bi-encoder encode query mỗi lần', 'Cache query embeddings cho repeated queries', 'Thấp'],
        ['Chunks quá dài hoặc quá ngắn', 'Overlap chunking 200-300 words, stride 100', 'Trung bình'],
    ]
    
    table = doc.add_table(rows=len(retrieval_opts), cols=3)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(retrieval_opts):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_heading('7.2. Tối ưu Gating Thresholds', level=2)
    doc.add_paragraph(
        'Ngưỡng gating cần được calibrate dựa trên dữ liệu thực tế:'
    )
    
    gating_opts = [
        'Chạy ScoreCalibrator trên ≥100 queries để học percentile distribution → auto-tune thresholds',
        'Rule 4 (legal keyword requirement) gây false ABSTAIN: '
        'ví dụ "Nguồn kinh phí..." bị abstain vì context thiếu "định nghĩa/là/được hiểu" '
        '→ Xem xét relax hoặc refine query type detection',
        'tier_local_min_margin=0.3 có thể quá cao nếu cross-encoder tạo score gần nhau '
        '→ Monitor phân bố margin thực tế và điều chỉnh',
        'Thêm Tier routing dựa trên query complexity (số từ, loại câu hỏi) ngoài score + margin',
        'A/B testing: so sánh chất lượng LOCAL vs API trên cùng câu hỏi, '
        'nếu LOCAL đạt ≥90% quality thì hạ threshold LOCAL',
    ]
    for opt in gating_opts:
        doc.add_paragraph(opt, style='List Bullet')
    
    doc.add_heading('7.3. Tối ưu LLM Generation', level=2)
    
    llm_opts = [
        ['Kỹ thuật', 'Mô tả', 'Kỳ vọng'],
        ['Prompt Engineering', 'Tối ưu prompt template cho pháp luật VN: few-shot examples, chain-of-thought', 'Tăng citation accuracy +15-20%'],
        ['Context compression', 'Dùng LLMLingua hoặc selective context để giảm tokens', 'Giảm chi phí API ~40%'],
        ['RAG-Fusion', 'Generate nhiều câu trả lời → merge/vote', 'Tăng accuracy, tăng latency'],
        ['Self-consistency', 'Generate 3 lần → majority vote', 'Tăng reliability'],
        ['Fine-tune LLM', 'Fine-tune Vistral/Vinallama trên QA pháp luật', 'Tăng chất lượng LOCAL tier'],
        ['Structured output', 'Force JSON output format (function calling)', 'Giảm parse errors'],
    ]
    
    table = doc.add_table(rows=len(llm_opts), cols=3)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(llm_opts):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_heading('7.4. Tối ưu chi phí (Cost Optimization)', level=2)
    
    doc.add_paragraph('Ước tính chi phí hiện tại (Gemini Flash):')
    
    add_styled_table(doc,
        ['Hạng mục', 'Ước tính', 'Ghi chú'],
        [
            ['Giá Gemini Flash', '$0.075 / 1M input tokens', 'gemini-2.0-flash-lite còn rẻ hơn'],
            ['Trung bình tokens/query', '~1500 tokens (context + prompt)', ''],
            ['1000 queries / ngày', '~$0.11 / ngày', 'Chưa tính output tokens'],
            ['Tier 1 LOCAL (miễn phí)', '~10-20% queries', 'Tiết kiệm $0.01-0.02/ngày'],
            ['Tier NONE (skip)', '~10% queries', 'Tiết kiệm hoàn toàn'],
            ['Tổng tiết kiệm nhờ 2-tier', '20-30% chi phí API', 'Có thể tăng khi tune LOCAL threshold'],
        ]
    )
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Chiến lược tối ưu chi phí:\n'
        '• Tăng tỷ lệ LOCAL bằng cách: hạ tier_local_min_score, hạ tier_local_min_margin\n'
        '• Cache kết quả: query giống nhau → trả về cache (Redis/SQLite)\n'
        '• Batch API calls: gom nhiều queries → 1 API call (giảm overhead)\n'
        '• Dùng Gemini Flash thay Gemini Pro (rẻ hơn ~10x, chất lượng gần bằng)\n'
        '• Monthly budget alert: set giới hạn API calls/ngày'
    )
    
    doc.add_heading('7.5. Tối ưu latency', level=2)
    
    latency_opts = [
        ['Bước', 'Latency hiện tại', 'Tối ưu', 'Mục tiêu'],
        ['FAISS Search', '~50-100ms', 'HNSW index thay Flat index', '<10ms'],
        ['Bi-encoder encode', '~100-200ms', 'ONNX Runtime / TensorRT', '<50ms'],
        ['Cross-encoder rerank', '~500-1000ms', 'Giảm candidates (top-30 thay top-50), batch_size=64', '<300ms'],
        ['Context build', '~5-10ms', 'Đã nhanh', '<5ms'],
        ['Gating', '<1ms', 'Đã nhanh', '<1ms'],
        ['LLM LOCAL', '~100-500ms', 'GPU offload, KV cache, flash attention', '<200ms'],
        ['LLM API', '~500-2000ms', 'Streaming (first token ~200ms)', '<500ms (first token)'],
        ['TỔNG', '~1.5-4s', 'Tối ưu toàn bộ', '<1s (LOCAL), <2s (API)'],
    ]
    
    table = doc.add_table(rows=len(latency_opts), cols=4)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(latency_opts):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 8. CÁC VẤN ĐỀ ĐÃ GIẢI QUYẾT
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('8. Các vấn đề đã giải quyết (Bug Fixes)', level=1)
    
    bugs = [
        ['#', 'Vấn đề', 'Nguyên nhân gốc', 'Cách sửa', 'Ảnh hưởng'],
        ['1', 'Metadata flattening bug (CRITICAL)', 
         '_load_components() chuyển {text, metadata} → {passage, **metadata} (flatten), '
         'nhưng retrieve_and_rerank đọc mapping.get("text","") → empty string',
         'Giữ nguyên raw JSON format {text, metadata} không flatten', 
         'Fix: chunk text không rỗng, keyword coverage > 0%, dedup hoạt động đúng'],
        ['2', 'Keyword coverage luôn 0%', 
         'Do bug #1: chunk.text rỗng → không có keyword nào match',
         'Fix bug #1 → text có nội dung → coverage ~90%', 
         'Gating không còn false ABSTAIN do thiếu keywords'],
        ['3', 'Chỉ 1 chunk sau dedup', 
         'Do bug #1: tất cả chunks có text rỗng → hash giống nhau → dedup loại hết',
         'Fix bug #1 → hash khác nhau → giữ đủ chunks', 
         '10 chunks retrieved thay vì 1'],
        ['4', 'tier_local_min_margin quá cao', 
         'Margin thực tế ~0.05-0.56, nhưng threshold ban đầu = 2.0',
         'Giảm tier_local_min_margin từ 2.0 → 0.3', 
         'Cho phép 10% queries route LOCAL'],
        ['5', 'GatingDecision thiếu tier/margin', 
         'Một số return points chưa có tier và margin fields',
         'Update tất cả 8+ return points trong evaluate()', 
         'Không còn AttributeError khi truy cập .tier'],
        ['6', 'Gemini API key not set', 
         '$env:GEMINI_API_KEY chưa được set trong terminal',
         'Cần set trước khi chạy: $env:GEMINI_API_KEY="..."', 
         'Gemini backend fallback sang placeholder'],
        ['7', 'Gemini model name deprecated', 
         'Auto-detect config set model_name="gemini-pro" (có thể deprecated)',
         'Cần update thành "gemini-2.0-flash" hoặc "gemini-1.5-flash"', 
         'Chưa fix — pending'],
    ]
    
    table = doc.add_table(rows=len(bugs), cols=5)
    table.style = 'Light Grid Accent 1'
    for row_idx, row_data in enumerate(bugs):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
                    if row_idx == 0:
                        run.bold = True
    
    doc.add_page_break()
    
    # ═══════════════════════════════════════════════════════════════════════════
    # 9. KẾ HOẠCH TIẾP THEO
    # ═══════════════════════════════════════════════════════════════════════════
    
    doc.add_heading('9. Kế hoạch tiếp theo (Roadmap)', level=1)
    
    doc.add_heading('Ưu tiên cao (nên làm ngay):', level=2)
    high_priority = [
        'Set GEMINI_API_KEY và test Tier 2 API thực tế với Gemini Flash',
        'Update gemini-pro → gemini-2.0-flash trong _auto_detect_config() (llm_client.py line ~122)',
        'Cài đặt LLM local thực tế (llama.cpp + Vistral-7B GGUF) cho Tier 1',
        'Chạy full evaluation 50-100 mẫu với Gemini API để đánh giá citation accuracy thực tế',
        'Relax Rule 4 (legal keyword requirement) — hiện gây false ABSTAIN cho một số queries hợp lệ',
    ]
    for item in high_priority:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Ưu tiên trung bình (tuần sau):', level=2)
    med_priority = [
        'Rebuild FAISS index sau khi dedup metadata (loại entries trùng)',
        'Implement BM25 hybrid retrieval (kết hợp sparse + dense)',
        'Tối ưu prompt template: thêm few-shot examples pháp luật VN',
        'Tạo goldset chất lượng cao hơn: manual annotation 100-200 queries',
        'Thêm caching layer (Redis/SQLite) cho repeated queries',
    ]
    for item in med_priority:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Ưu tiên thấp (tháng sau):', level=2)
    low_priority = [
        'Deploy FastAPI backend + React frontend',
        'Implement streaming cho Tier 2 API',
        'Fine-tune LLM local trên dữ liệu QA pháp luật',
        'A/B testing chất lượng LOCAL vs API',
        'Scale dữ liệu: thêm Luật, Nghị quyết, Công văn',
        'Self-consistency voting + RAG-Fusion',
        'Chuyển FAISS sang HNSW index khi >10K vectors',
        'ONNX Runtime cho bi-encoder inference',
    ]
    for item in low_priority:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # ─── Footer note ───
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 50)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Báo cáo được tạo tự động ngày {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    # ═══════════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════════
    
    output_path = os.path.join(os.path.dirname(__file__), 
                               'cross-encoder', 'generation', 'outputs', 
                               'BaoCao_LegalRAG_2Tier.docx')
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ Báo cáo đã được tạo: {output_path}")
    print(f"   Kích thước: {os.path.getsize(output_path):,} bytes")
    
    return output_path


if __name__ == "__main__":
    create_report()
